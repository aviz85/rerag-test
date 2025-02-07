import streamlit as st
import asyncio
from reag_client import ReagClient, Document, OPENROUTER_MODELS
import io
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

def read_docx(file):
    doc = DocxDocument(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_pdf(file):
    pdf = PdfReader(file)
    return "\n".join([page.extract_text() for page in pdf.pages])

st.set_page_config(page_title="ReAG Demo", page_icon="🎓", layout="wide")

st.title("🎓 ReAG Demo")
st.markdown("### Reasoning Augmented Generation Demo")

# Sidebar configuration
with st.sidebar:
    st.header("Configuration")
    model = st.selectbox(
        "Select Model",
        OPENROUTER_MODELS,
        index=1  # Default to Claude 3.5 Sonnet
    )
    api_key = st.text_input("OpenRouter API Key (Required)", type="password")
    system_prompt = st.text_area("System Prompt (Optional)", "")

    if not api_key:
        st.warning("Please enter your OpenRouter API key to use the demo. Get one at https://openrouter.ai/keys")

# Main content
col1, col2 = st.columns(2)

with col1:
    st.subheader("Document Input")
    
    # Add file upload
    uploaded_file = st.file_uploader("Upload Document", type=['txt', 'md', 'py', 'js', 'html', 'css', 'json', 'yaml', 'yml', 'docx', 'pdf'])
    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1].lower()
        try:
            if file_type == 'pdf':
                content = read_pdf(uploaded_file)
            elif file_type == 'docx':
                content = read_docx(uploaded_file)
            else:
                content = uploaded_file.getvalue().decode()
            doc_name = uploaded_file.name
            doc_content = content
        except Exception as e:
            st.error(f"Error reading file: {str(e)}")
            doc_name = "Example Doc"
            doc_content = "Enter your document content here..."
    else:
        doc_name = st.text_input("Document Name", "Example Doc")
        doc_content = st.text_area("Document Content", "Enter your document content here...")
    
    doc_url = st.text_input("URL (Optional)", "")
    doc_source = st.text_input("Source (Optional)", "")
    
    if st.button("Add Document"):
        if "documents" not in st.session_state:
            st.session_state.documents = []
        
        doc = Document(
            name=doc_name,
            content=doc_content,
            metadata={
                "url": doc_url,
                "source": doc_source,
            }
        )
        st.session_state.documents.append(doc)
        st.success("Document added!")

with col2:
    st.subheader("Query")
    query = st.text_area("Enter your query", "")
    
    if st.button("Submit Query"):
        if "documents" not in st.session_state or not st.session_state.documents:
            st.error("Please add at least one document first!")
        else:
            with st.spinner("Processing query..."):
                async def run_query():
                    async with ReagClient(
                        model=model,
                        system=system_prompt if system_prompt else None,
                        api_key=api_key
                    ) as client:
                        response = await client.query(
                            query,
                            documents=st.session_state.documents
                        )
                        return response

                response = asyncio.run(run_query())
                
                st.subheader("Response")
                st.write("Content:", response.content)
                st.write("Reasoning:", response.reasoning)
                st.write("Is Irrelevant:", response.is_irrelevant)
                st.write("Document Used:", response.document.name)

# Display current documents
if "documents" in st.session_state and st.session_state.documents:
    st.subheader("Current Documents")
    for i, doc in enumerate(st.session_state.documents):
        with st.expander(f"Document {i+1}: {doc.name}"):
            st.write("Content:", doc.content)
            st.write("Metadata:", doc.metadata)
            if st.button(f"Remove Document {i+1}"):
                st.session_state.documents.pop(i)
                st.rerun() 